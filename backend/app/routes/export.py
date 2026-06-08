from __future__ import annotations

import io
import logging
from uuid import UUID

from fastapi import APIRouter, Depends
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.core.auth import get_current_user
from app.core.dependencies import get_db
from app.core.entitlements import effective_plan, require_feature
from app.core.exceptions import AppError
from app.db.models import Conversation, ConversationTurn, User
from app.routes.workspaces import _assert_workspace_access

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    conversation_id: UUID
    format: str  # "pdf" | "docx"


def _build_pdf(title: str, turns: list) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, title[:80], new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=10)
    pdf.ln(4)
    for turn in turns:
        pdf.set_font("Helvetica", "B", 10)
        pdf.multi_cell(0, 6, f"Q: {turn.question or ''}")
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 6, f"A: {turn.answer or ''}")
        pdf.ln(3)
    return bytes(pdf.output())


def _build_docx(title: str, turns: list) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading(title[:80], 0)
    for turn in turns:
        doc.add_heading("Question", level=2)
        doc.add_paragraph(turn.question or "")
        doc.add_heading("Answer", level=2)
        doc.add_paragraph(turn.answer or "")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/conversation")
def export_conversation(
    body: ExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> Response:
    # Step 1: gate — free users blocked
    require_feature(current_user, "export")

    # Step 2: docx requires business plan
    fmt = body.format.lower()
    if fmt not in ("pdf", "docx"):
        raise AppError(400, "invalid_format", "format must be 'pdf' or 'docx'.")
    if fmt == "docx" and effective_plan(current_user) != "business":
        raise AppError(403, "docx_requires_business", "Word export requires Business plan.")

    # Step 3: load conversation and verify access
    conversation = (
        db.query(Conversation)
        .filter(
            Conversation.id == body.conversation_id,
            Conversation.is_active.is_(True),
        )
        .first()
    )
    if conversation is None:
        raise AppError(404, "conversation_not_found", "Conversation not found.")

    if not conversation.workspace_id:
        raise AppError(404, "conversation_not_found", "Conversation not found.")
    _assert_workspace_access(conversation.workspace_id, current_user, db)

    # Step 4: load turns ordered by turn_index
    turns = (
        db.query(ConversationTurn)
        .filter(ConversationTurn.conversation_id == body.conversation_id)
        .order_by(ConversationTurn.turn_index)
        .all()
    )

    title = conversation.title or "Conversation"

    # Step 5: generate file
    if fmt == "pdf":
        file_bytes = _build_pdf(title, turns)
        media_type = "application/pdf"
        filename = "conversation.pdf"
    else:
        file_bytes = _build_docx(title, turns)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = "conversation.docx"

    # Step 6: return file response
    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
